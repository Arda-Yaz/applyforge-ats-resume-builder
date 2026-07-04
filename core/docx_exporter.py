from pathlib import Path

from docx import Document as create_document
from docx.document import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Any, cast

def _set_document_margins(document: DocxDocument) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)


def _configure_styles(document: DocxDocument) -> None:
    styles = document.styles

    normal = cast(Any, styles["Normal"])
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = cast(Any, styles[style_name])
        style.font.name = "Arial"
        style.font.bold = True

    cast(Any, styles["Heading 1"]).font.size = Pt(14)
    cast(Any, styles["Heading 2"]).font.size = Pt(11)
    cast(Any, styles["Heading 3"]).font.size = Pt(10)

def _add_section_heading(document: DocxDocument, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)

    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(2)


def _add_bullet(document: DocxDocument, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)

    run = paragraph.add_run(text)
    run.font.size = Pt(10)


def _add_role_header(
    document: DocxDocument,
    title: str,
    organization: str,
    location: str,
    start: str,
    end: str
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)

    left = paragraph.add_run(f"{title} — {organization}")
    left.bold = True
    left.font.size = Pt(10)

    right = paragraph.add_run(f" | {location} | {start} - {end}")
    right.font.size = Pt(10)


def _add_project_header(document: DocxDocument, name: str, tech: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)

    title = paragraph.add_run(name)
    title.bold = True
    title.font.size = Pt(10)

    if tech:
        tech_run = paragraph.add_run(f" | {', '.join(tech)}")
        tech_run.italic = True
        tech_run.font.size = Pt(10)


def export_resume_to_docx(
    profile: dict,
    grouped_resume: dict,
    target_title: str,
    output_path: str = "exports/generated_resume.docx"
) -> str:
    document = create_document()

    _set_document_margins(document)
    _configure_styles(document)

    personal = profile["personal"]

    # Header
    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name_run = name_paragraph.add_run(personal["name"])
    name_run.bold = True
    name_run.font.size = Pt(16)

    contact_paragraph = document.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph.paragraph_format.space_after = Pt(4)

    contact_text = (
        f"{personal['location']} | {personal['email']} | {personal['phone']} | "
        f"{personal['linkedin']} | {personal['github']}"
    )

    contact_run = contact_paragraph.add_run(contact_text)
    contact_run.font.size = Pt(9)

    # Target
    target_paragraph = document.add_paragraph()
    target_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target_paragraph.paragraph_format.space_after = Pt(6)

    target_run = target_paragraph.add_run(target_title)
    target_run.bold = True
    target_run.font.size = Pt(10)

    # Summary
    _add_section_heading(document, "Summary")
    summary_paragraph = document.add_paragraph()
    summary_paragraph.paragraph_format.space_after = Pt(2)
    summary_run = summary_paragraph.add_run(profile["summary"])
    summary_run.font.size = Pt(10)

    # Technical Skills
    _add_section_heading(document, "Technical Skills")
    for category, skills in profile.get("skills", {}).items():
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)

        category_run = paragraph.add_run(f"{category}: ")
        category_run.bold = True
        category_run.font.size = Pt(10)

        skills_run = paragraph.add_run(", ".join(skills))
        skills_run.font.size = Pt(10)

    # Experience
    if grouped_resume.get("experience"):
        _add_section_heading(document, "Experience")

        for exp in grouped_resume["experience"]:
            _add_role_header(
                document=document,
                title=exp["title"],
                organization=exp["company"],
                location=exp["location"],
                start=exp["start"],
                end=exp["end"]
            )

            for bullet in exp.get("selected_bullets", []):
                _add_bullet(document, bullet["text"])

    # Projects
    if grouped_resume.get("projects"):
        _add_section_heading(document, "Projects")

        for project in grouped_resume["projects"]:
            _add_project_header(
                document=document,
                name=project["name"],
                tech=project.get("tech", [])
            )

            for bullet in project.get("selected_bullets", []):
                _add_bullet(document, bullet["text"])

    # Education
    _add_section_heading(document, "Education")
    for edu in profile.get("education", []):
        _add_role_header(
            document=document,
            title=edu["degree"],
            organization=edu["school"],
            location=edu["location"],
            start=edu["start"],
            end=edu["end"]
        )

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    document.save(str(output))

    return str(output)