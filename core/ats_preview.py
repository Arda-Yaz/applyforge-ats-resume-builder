from pathlib import Path

from docx import Document


SUSPICIOUS_CHARACTERS = [
    "�",
    "¨",
    "˘",
    "¸",
    "˙",
    "˜",
]


REQUIRED_SECTIONS = [
    "SUMMARY",
    "TECHNICAL SKILLS",
    "EXPERIENCE",
    "PROJECTS",
    "EDUCATION",
]


def extract_docx_text(docx_path: str) -> str:
    path = Path(docx_path)

    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    document = Document(str(path))
    text_parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


def detect_suspicious_characters(parsed_text: str) -> list[str]:
    found = []

    for character in SUSPICIOUS_CHARACTERS:
        if character in parsed_text:
            found.append(character)

    return found


def check_required_sections(parsed_text: str) -> dict:
    normalized_text = parsed_text.upper()

    return {
        section: section in normalized_text
        for section in REQUIRED_SECTIONS
    }


def count_keyword_matches(parsed_text: str, job_keywords: list[str]) -> dict:
    normalized_text = parsed_text.lower()

    matched = []
    missing = []

    for keyword in job_keywords:
        keyword_normalized = keyword.lower()

        if keyword_normalized in normalized_text:
            matched.append(keyword)
        else:
            missing.append(keyword)

    return {
        "matched": matched,
        "missing": missing,
    }


def calculate_ats_parse_score(
    parsed_text: str,
    section_status: dict,
    suspicious_characters: list[str],
    keyword_status: dict,
) -> int:
    score = 100

    missing_sections = [
        section for section, exists in section_status.items()
        if not exists
    ]

    score -= len(missing_sections) * 10
    score -= len(suspicious_characters) * 10

    if len(parsed_text.strip()) < 500:
        score -= 15

    total_keywords = len(keyword_status["matched"]) + len(keyword_status["missing"])

    if total_keywords > 0:
        keyword_match_ratio = len(keyword_status["matched"]) / total_keywords

        if keyword_match_ratio < 0.4:
            score -= 15
        elif keyword_match_ratio < 0.6:
            score -= 8

    return max(0, min(100, score))


def build_ats_preview_report(parsed_text: str, job_keywords: list[str]) -> dict:
    suspicious_characters = detect_suspicious_characters(parsed_text)
    section_status = check_required_sections(parsed_text)
    keyword_status = count_keyword_matches(parsed_text, job_keywords)

    parse_score = calculate_ats_parse_score(
        parsed_text=parsed_text,
        section_status=section_status,
        suspicious_characters=suspicious_characters,
        keyword_status=keyword_status,
    )

    return {
        "parse_score": parse_score,
        "parsed_text": parsed_text,
        "suspicious_characters": suspicious_characters,
        "section_status": section_status,
        "keyword_status": keyword_status,
        "character_count": len(parsed_text),
        "word_count": len(parsed_text.split()),
    }